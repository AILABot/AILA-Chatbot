"""
Evidence Ingestion & Report Generation (Uploads → Prompts → DOCX)
==================================================================

Purpose
-------
Utilities for persisting user uploads, extracting evidence (images, audio, PDFs, text),
building multimodal LangChain messages (text + images), and generating a DOCX report
(«ΜΗΝΥΣΗ») with embedded previews and excerpts.

Key Functions
-------------
- guess_ext           : Infer file extension from a filename.
- persist_upload      : Save UploadFile to disk, return FileRec metadata.
- to_data_url         : Convert a file to a base64 data URL for inline model consumption.
- transcribe_audio    : Transcribe audio files via OpenAI Whisper.
- extract_text_from_pdf: Extract plain text from all PDF pages.
- safe_read_text      : Read text files robustly (UTF-8 with ignore errors).
- build_evidence      : Produce human-readable evidence lines from uploaded files.
- build_messages      : Build LangChain HumanMessage payload with inline media and evidence.
- create_word_file    : Generate a Greek-labeled DOCX («ΜΗΝΥΣΗ») with previews of attachments.

Dependencies
------------
FastAPI (UploadFile), pypdf, python-docx, OpenAI SDK, LangChain, and your FileRec model.
Requires settings.API_KEY for Whisper transcription.
"""

import base64
from openai import OpenAI
from backend.database.config.config import settings
from pypdf import PdfReader
from fastapi import UploadFile
from langchain_core.messages import HumanMessage
from backend.api.models import FileRec
import uuid,os,shutil
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from datetime import date
import uuid 

def guess_ext(filename: str) -> str:
    """
    Extract the file extension from a filename.

    Args:
        filename (str): Input filename.

    Returns:
        str: Lowercased file extension (e.g., ".pdf").
    """
    _, ext = os.path.splitext(filename or "")
    return ext.lower()

def persist_upload(f: UploadFile) -> FileRec:
    """
    Save an uploaded file to the server.

    - Generates a unique filename using UUID.
    - Preserves original extension.
    - Stores file under 'uploads/' directory.
    - Returns metadata as FileRec.

    Args:
        f (UploadFile): The file uploaded by the client.

    Returns:
        FileRec: Metadata containing original name, storage path, and MIME type.
    """
    ext = guess_ext(f.filename)
    new_name = f"{uuid.uuid4().hex}{ext}"
    dest = os.path.join('uploads', new_name)
    with open(dest, "wb") as out:
        shutil.copyfileobj(f.file, out)
    return FileRec(original=f.filename or new_name, path=dest, mime=(f.content_type or "").lower())


def normalize_mime(mt: str) -> str:
    """
    Keep only 'image/<subtype>' and map oddities (jpg -> jpeg).
    Strip any extra parameters after ';'.
    """
    if not mt:
        return "image/png"  # sane default
    core = mt.split(";")[0].strip().lower()
    if core == "image/jpg":
        core = "image/jpeg"
    # allow only common types (tighten if needed)
    allowed = {"image/png", "image/jpeg", "image/webp", "image/gif"}
    return core if core in allowed else "image/png"

def to_data_url(path: str, mime: str) -> str:
    """
    Convert a file into a base64-encoded data URL.

    Args:
        path (str): Path to file.
        mime (str): MIME type of file.

    Returns:
        str: Data URL string (data:{mime};base64,...).
    """
    mime = normalize_mime(mime)
    with open(path, 'rb') as f:
        b64 = base64.b64encode(f.read()).decode('utf-8')
    return f"data:{mime};base64,{b64}"


def transcribe_audio(path:str) -> str:
    """
    Transcribe an audio file using OpenAI Whisper.

    Args:
        path (str): Path to audio file.

    Returns:
        str: Transcribed text.
    """
    oai = OpenAI(api_key = settings.API_KEY)
    with open(path,"rb") as f:
        tr = oai.audio.transcriptions.create(model='whisper-1',file=f)
    return getattr(tr,'text',str(tr))

def extract_text_from_pdf(path: str) -> str:
    """
    Extract plain text from all pages of a PDF.

    Args:
        path (str): Path to PDF file.

    Returns:
        str: Concatenated text from all pages.
    """
    reader = PdfReader(path)
    parts = []
    for p in reader.pages:
        txt = p.extract_text() or ""
        parts.append(txt)
    return "\n".join(parts)

def safe_read_text(path:str) -> str:
    """
    Read text from a file safely (UTF-8 with fallback).

    Args:
        path (str): Path to text file.

    Returns:
        str: File contents as string.
    """
    txt = open(path,'r',encoding='utf-8',errors='ignore').read()
    return txt

def build_evidence(files):
    """
    Build evidence strings from uploaded files for inclusion in prompts or reports.

    Supports:
        - Images: listed only by filename.
        - Audio: transcribed to text.
        - PDFs: text extracted.
        - Plain text / CSV: contents read.
        - Others: ignored.

    Args:
        files (list[FileRec]): List of uploaded file metadata.

    Returns:
        list[str]: Evidence descriptions (with text or transcription where applicable).
    """
    evidence_lines = []
    for f in files:
        mt = (f.mime or "").lower()
        if mt.startswith('image/'):
            evidence_lines.append(f"-Image: {f.original}")
        elif mt.startswith('audio/'):
            transcript = transcribe_audio(f.path)
            evidence_lines.append(f" Sound {f.original}")

        elif mt == 'application/pdf':
            txt = extract_text_from_pdf(f.path)
            evidence_lines.append(f"- PDF: {f.original}\n Text:\n {txt}")

        elif mt in ("text/plain","text/csv") or f.original.lower().endswith(".csv"):
            snippet = safe_read_text(f.path)
            evidence_lines.append(f"- Text: {f.original}\n Text:\n {snippet}")
        
    return evidence_lines

def build_messages(prompt_text:str, files: list[UploadFile]):
    """
    Build a chat message (LangChain HumanMessage) including user input and file evidence.

    - Converts images into inline data URLs for model consumption.
    - Transcribes audio files.
    - Extracts text from PDFs and plain text files.
    - Appends evidence summary to prompt.

    Args:
        prompt_text (str): The main user prompt/question.
        files (list[UploadFile]): Uploaded files associated with the prompt.

    Returns:
        list[HumanMessage]: LangChain-formatted messages containing text + media.
    """
    parts = [{'type':'text','text':(prompt_text)}]
    evidence_lines = []

    for f in files:
        mt = (f.mime or "").lower()
        if mt.startswith('image/'):
            img_url = f.public_url or to_data_url(f.path,mt)
            parts.append({'type':'image_url','image_url':{"url": img_url}})
            evidence_lines.append(f"-Image: {f.original}")

        elif mt.startswith('audio/'):
            transcript = transcribe_audio(f.path)
            evidence_lines.append(f" Sound {f.original} \n Recording:\n {transcript}")

        elif mt == 'application/pdf':
            txt = extract_text_from_pdf(f.path)
            evidence_lines.append(f"- PDF: {f.original}\n Text:\n {txt}")

        elif mt in ("text/plain","text/csv") or f.original.lower().endswith(".csv"):
            snippet = safe_read_text(f.path)
            evidence_lines.append(f"- Text: {f.original}\n Text:\n {snippet}")

        else:
            evidence_lines.append(f"- No supported type. {f.original} will be ignored")
        
    if evidence_lines:
        parts[0]['text'] += "\n\n Evidence from Uploads\n" + "\n".join(evidence_lines)

    return [HumanMessage(content=parts)]




def create_word_file(text:str, files:list, output_path:str = 'uploads'):
    """
    Generate a DOCX report containing:
        - Title ("ΜΗΝΥΣΗ")
        - Date
        - User-provided body text
        - Evidence section:
            * Images embedded directly
            * Text, PDF, DOCX previews included
            * Other file types listed only

    Args:
        text (str): Main body text of the complaint/report.
        files (list): Files to be attached and previewed.
        output_path (str, optional): Directory to save the DOCX. Defaults to 'uploads'.

    Returns:
        tuple(str, str): Output directory path, generated filename.
    """
    filename = f"{uuid.uuid4()}.docx"
    out_path = os.path.join(output_path, filename)
    doc = Document()

    def _add_header(doc, text: str):
        p = doc.add_paragraph(text)
        if p.runs:
            p.runs[0].bold = True

    def _add_text_block(doc, text: str, max_lines: int = 100):
        lines = (text or "").splitlines()
        for line in lines[:max_lines]:
            doc.add_paragraph(line.strip())
        if len(lines) > max_lines:
            doc.add_paragraph("... [περικοπή περιεχομένου] ...")



    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    # Compute once (fits images to margins)
    max_pic_width = section.page_width - section.left_margin - section.right_margin

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")

     # Title block
    doc.add_paragraph("ΜΗΝΥΣΗ").runs[0].bold = True
    doc.paragraphs[-1].runs[0].font.size = Pt(16)

    # Date & place line
    doc.add_paragraph(f"Ημερομηνία: {date.today().strftime('%d/%m/%Y')}")

    # Body: preserve your model’s layout
    for line in text.split("\n"):
        doc.add_paragraph(line)

    # Evidence section (images embedded; others listed)
    doc.add_paragraph("")  # spacer
    h = doc.add_paragraph("VI. Συνημμένα")
    h.runs[0].bold = True

    idx = 1
    for f in files or []:
        mime = (getattr(f, "mime", "") or "").lower()
        path = getattr(f, "path", None)
        name = getattr(f, "original", os.path.basename(getattr(f, "path", "αρχείο")))

        # File listing line
        doc.add_paragraph(f"{idx}) {name} — {mime if mime else 'άγνωστος τύπος'}")

        if path and os.path.exists(path):
            # ----- Images -----
            if mime.startswith("image/"):
                try:
                    doc.add_picture(path, width=max_pic_width)
                except Exception:
                    doc.add_paragraph("[Σφάλμα: Αδυναμία ενσωμάτωσης εικόνας]")

            # ----- Plain text -----
            elif mime == "text/plain" or (mime == "" and path.lower().endswith(".txt")):
                try:
                    with open(path, "r", encoding="utf-8", errors="ignore") as ftxt:
                        contents = ftxt.read()
                    _add_header(doc, "---- Περιεχόμενο αρχείου ----")
                    _add_text_block(doc, contents, max_lines=100)
                    _add_header(doc, "---- Τέλος περιεχομένου ----")
                except Exception:
                    doc.add_paragraph("[Σφάλμα: Αδυναμία ανάγνωσης αρχείου κειμένου]")

            # ----- PDF preview (first ~5 pages) -----
            elif mime == "application/pdf" or path.lower().endswith(".pdf"):
                if PdfReader is None:
                    doc.add_paragraph("[Σημείωση: Εγκαταστήστε PyPDF2 για προεπισκόπηση PDF]")
                else:
                    try:
                        reader = PdfReader(path)
                        pages_to_read = min(5, len(reader.pages))
                        extracted = []
                        for i in range(pages_to_read):
                            try:
                                txt = reader.pages[i].extract_text() or ""
                            except Exception:
                                txt = ""
                            extracted.append(txt)
                        preview = "\n".join(extracted).strip()
                        if preview:
                            _add_header(doc, "---- Απόσπασμα PDF ----")
                            _add_text_block(doc, preview, max_lines=120)
                            _add_header(doc, "---- Τέλος αποσπάσματος ----")
                        else:
                            doc.add_paragraph("[Δεν βρέθηκε εξαγώγιμο κείμενο στο PDF]")
                    except Exception:
                        doc.add_paragraph("[Σφάλμα: Αδυναμία ανάγνωσης PDF]")

            # ----- DOCX preview (first ~100 paragraphs) -----
            elif (
                mime
                in (
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    "application/msword",
                )
                or path.lower().endswith(".docx")
            ):
                try:
                    d = Document(path)
                    paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
                    preview = "\n".join(paras[:100])
                    if preview:
                        _add_header(doc, "---- Απόσπασμα DOCX ----")
                        _add_text_block(doc, preview, max_lines=120)
                        _add_header(doc, "---- Τέλος αποσπάσματος ----")
                    else:
                        doc.add_paragraph("[Κενό ή μη εξαγώγιμο περιεχόμενο DOCX]")
                except Exception:
                    doc.add_paragraph("[Σφάλμα: Αδυναμία ανάγνωσης DOCX]")

            # ----- Other/unknown types: just listed above -----
            # (You can add more handlers here later)
        else:
            doc.add_paragraph("[Προειδοποίηση: Το μονοπάτι δεν υπάρχει ή δεν δόθηκε]")

        idx += 1


    doc.save(out_path)
    return output_path,filename